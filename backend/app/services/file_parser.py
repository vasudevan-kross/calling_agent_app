import pandas as pd
import PyPDF2
import pdfplumber
from docx import Document
from typing import List, Dict, Any
import io
import re


class FileParserService:
    """Service for parsing various file formats to extract lead data"""

    def __init__(self):
        self.phone_pattern = re.compile(r'\+?1?\d{9,15}')
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')

    async def parse_file(
        self,
        filename: str,
        content: bytes
    ) -> List[Dict[str, Any]]:
        """Parse file and extract contact data"""
        file_ext = filename.split(".")[-1].lower()

        if file_ext in ["xlsx", "xls"]:
            return await self._parse_excel(content)
        elif file_ext == "csv":
            return await self._parse_csv(content)
        elif file_ext == "pdf":
            return await self._parse_pdf(content)
        elif file_ext in ["docx", "doc"]:
            return await self._parse_docx(content)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    async def _parse_excel(self, content: bytes) -> List[Dict[str, Any]]:
        """Parse Excel file"""
        df = pd.read_excel(io.BytesIO(content))
        return self._dataframe_to_leads(df)

    async def _parse_csv(self, content: bytes) -> List[Dict[str, Any]]:
        """Parse CSV file"""
        df = pd.read_csv(io.BytesIO(content))
        return self._dataframe_to_leads(df)

    async def _parse_pdf(self, content: bytes) -> List[Dict[str, Any]]:
        """Parse PDF file"""
        leads = []

        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        # Extract contacts from text
                        contacts = self._extract_contacts_from_text(text)
                        leads.extend(contacts)

                    # Try to extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            table_leads = self._dataframe_to_leads(df)
                            leads.extend(table_leads)
        except Exception as e:
            # Fallback to PyPDF2
            print(f"pdfplumber failed: {e}, trying PyPDF2")
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    contacts = self._extract_contacts_from_text(text)
                    leads.extend(contacts)

        return leads

    async def _parse_docx(self, content: bytes) -> List[Dict[str, Any]]:
        """Parse Word document"""
        doc = Document(io.BytesIO(content))
        leads = []

        # Try to extract from tables first
        for table in doc.tables:
            df = self._table_to_dataframe(table)
            if not df.empty:
                table_leads = self._dataframe_to_leads(df)
                leads.extend(table_leads)

        # If no tables or no leads found, extract from paragraphs
        if not leads:
            text = "\n".join([para.text for para in doc.paragraphs])
            leads = self._extract_contacts_from_text(text)

        return leads

    def _dataframe_to_leads(self, df: pd.DataFrame) -> List[Dict[str, Any]]:
        """Convert DataFrame to list of lead dictionaries"""
        # Normalize column names
        column_mapping = {
            "name": ["name", "contact_name", "full_name", "person", "contact"],
            "business_name": ["business", "company", "business_name", "organization", "org"],
            "phone": ["phone", "telephone", "mobile", "contact_number", "tel", "cell"],
            "email": ["email", "email_address", "e-mail", "mail"],
            "address": ["address", "street", "location", "addr"],
            "city": ["city", "town"],
            "state": ["state", "province", "region"],
            "country": ["country"],
        }

        normalized_df = pd.DataFrame()
        for target_col, possible_names in column_mapping.items():
            for col in df.columns:
                if str(col).lower().strip() in possible_names:
                    normalized_df[target_col] = df[col]
                    break

        # Ensure phone exists
        if "phone" not in normalized_df.columns:
            # Try to find phone numbers in any column
            for col in df.columns:
                phone_found = False
                for val in df[col].dropna():
                    if self.phone_pattern.search(str(val)):
                        normalized_df["phone"] = df[col]
                        phone_found = True
                        break
                if phone_found:
                    break

        if "phone" not in normalized_df.columns:
            raise ValueError("Could not find phone number column in file")

        # Convert to lead dictionaries
        leads = []
        for _, row in normalized_df.iterrows():
            lead = {
                "name": str(row.get("name", "")).strip() or "Unknown",
                "business_name": str(row.get("business_name", "")).strip() or None,
                "phone": str(row.get("phone", "")).strip(),
                "email": str(row.get("email", "")).strip() or None,
                "address": str(row.get("address", "")).strip() or None,
                "city": str(row.get("city", "")).strip() or None,
                "state": str(row.get("state", "")).strip() or None,
                "country": str(row.get("country", "")).strip() or None,
                "source": "file_import"
            }

            # Clean up None values that are string "nan"
            for key, value in lead.items():
                if value and isinstance(value, str) and value.lower() in ["nan", "none", ""]:
                    lead[key] = None

            # Only add if phone is valid
            if lead["phone"] and lead["phone"] not in ["nan", "None", ""]:
                leads.append(lead)

        return leads

    def _extract_contacts_from_text(self, text: str) -> List[Dict[str, Any]]:
        """Extract contacts from unstructured text"""
        lines = text.split("\n")
        leads = []

        for i, line in enumerate(lines):
            # Find phone numbers in line
            phones = self.phone_pattern.findall(line)
            if phones:
                # Get context (surrounding lines)
                context_start = max(0, i - 2)
                context_end = min(len(lines), i + 3)
                context = " ".join(lines[context_start:context_end])

                # Find email in context
                emails = self.email_pattern.findall(context)

                # Use line or context as name
                name = line[:100].strip() if line.strip() else "Unknown Contact"

                lead = {
                    "name": name,
                    "phone": phones[0],
                    "email": emails[0] if emails else None,
                    "source": "file_import"
                }
                leads.append(lead)

        return leads

    def _table_to_dataframe(self, table) -> pd.DataFrame:
        """Convert Word table to DataFrame"""
        data = []
        keys = None

        for i, row in enumerate(table.rows):
            text = [cell.text.strip() for cell in row.cells]
            if i == 0:
                keys = text
            else:
                data.append(text)

        if keys and data:
            return pd.DataFrame(data, columns=keys)
        return pd.DataFrame()
